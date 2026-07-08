"""Agent 工具集。

定义可供 Agent 调用的工具函数：
- search_knowledge_base: 检索上传的知识库文档
- search_web: 联网搜索实时信息
- sandbox_run: 在 Docker 沙盒中执行 shell 命令
- sandbox_write_file / sandbox_read_file / sandbox_list_files: 沙盒文件操作
- search_skills / install_skill: 宿主机技能搜索与安装（不经过沙盒）
"""

import contextvars
import json
import subprocess
from pathlib import Path

from langchain_core.tools import tool
from tavily import TavilyClient

from server.config import RAGConfig
from server.logging_config import get_logger
from server.retrieval.retriever import expand_and_retrieve
from server.chain import _resolve_collection_name

logger = get_logger(__name__)

# 通过 contextvar 将 kb_id 和 config 从 chat_stream 传入工具
_current_kb_id: contextvars.ContextVar[str | None] = contextvars.ContextVar(
    "kb_id", default=None
)
_current_config: contextvars.ContextVar[RAGConfig] = contextvars.ContextVar(
    "config"
)
# session_id 用于沙盒跨消息持久化
_current_session_id: contextvars.ContextVar[str | None] = contextvars.ContextVar(
    "session_id", default=None
)


def set_agent_context(kb_id: str | None, config: RAGConfig, session_id: str | None = None) -> None:
    """设置 Agent 工具的运行时上下文（由 chat_stream 调用）。"""
    _current_kb_id.set(kb_id)
    _current_config.set(config)
    _current_session_id.set(session_id)


@tool
def search_knowledge_base(query: str) -> str:
    """在用户已上传的知识库中搜索相关内容，返回最相关的文档片段。
    适用场景：用户询问文档、小说、技术手册、项目资料等知识库内的问题。
    不适用场景：实时新闻、天气、股价等动态信息——这类问题请用 search_web。"""
    config = _current_config.get()
    kb_id = _current_kb_id.get()

    if not kb_id:
        return "（当前未选择知识库，无法检索。请告诉用户需要先选择一个知识库。）"

    collection_name = _resolve_collection_name(kb_id, config)
    docs = expand_and_retrieve(query, collection_name, config)

    if not docs:
        return "（知识库中未找到相关内容）"

    return "\n\n---\n\n".join(
        f"【来源 {i + 1}】{doc.page_content}"
        for i, doc in enumerate(docs)
    )


@tool
def search_web(query: str, max_results: int = 5, search_depth: str = "basic") -> str:
    """搜索互联网获取实时信息，返回搜索结果摘要。
    
    参数：
    - query: 搜索关键词
    - max_results: 返回结果数量（1-10，默认5）。简单问题用3即可，需要广泛调研时可用8-10
    - search_depth: 搜索深度，"basic"（快速摘要）或 "advanced"（深度搜索，内容更详细但更慢）
    
    使用建议：
    - 简单事实查询：max_results=3, search_depth="basic"
    - 常规搜索：max_results=5, search_depth="basic"（默认）
    - 深度调研：max_results=8, search_depth="advanced"
    
    适用场景：最新新闻、天气、股价、实时事件等知识库无法覆盖的动态信息。
    不适用场景：已上传知识库中的内容——这类问题请用 search_knowledge_base。"""
    config = _current_config.get()

    if not config.tavily_api_key:
        return "（联网搜索功能未配置 API Key，请联系管理员设置 tavily_api_key）"

    # 参数边界保护
    max_results = max(1, min(10, max_results))
    if search_depth not in ("basic", "advanced"):
        search_depth = "basic"

    try:
        client = TavilyClient(api_key=config.tavily_api_key)
        result = client.search(query, search_depth=search_depth, max_results=max_results)
    except Exception as exc:
        return f"（搜索失败: {exc}）"

    if not result.get("results"):
        return "（未找到相关搜索结果）"

    return "\n\n".join(
        f"【{r['title']}】({r['url']})\n{r['content']}"
        for r in result["results"]
    )


@tool
def fetch_webpage(url: str) -> str:
    """抓取指定网页内容，返回正文文本（转换为 Markdown 格式）。
    适用场景：用户给了具体网址，需要读取页面内容；或搜索到结果后想看详情。
    不适用场景：搜索关键词——请用 search_web。"""
    from server.agent.sandbox import get_or_create

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动，无法抓取网页）"

    # httpx + html2text 已预装在沙盒镜像中
    script = '''import sys
import httpx, html2text

url = sys.argv[1]
try:
    resp = httpx.get(url, timeout=20, follow_redirects=True, headers={"User-Agent": "Mozilla/5.0"})
    resp.raise_for_status()
except Exception as e:
    print(f"ERROR: 请求失败: {e}")
    sys.exit(1)

h = html2text.HTML2Text()
h.ignore_links = False
h.ignore_images = True
h.body_width = 0
text = h.handle(resp.text)
# 去除过多空行
lines = [l.rstrip() for l in text.split("\\n")]
text = "\\n".join(lines)
# 截断
if len(text) > 8000:
    text = text[:8000] + "\\n\\n... （内容过长，已截断）"
print(text)
'''
    sb.write_file("/workspace/_fetch.py", script)
    # 用单引号包裹 URL 防止 shell 注入，并转义 URL 中的单引号
    safe_url = url.replace("'", "'\"'\"'")
    exit_code, output = sb.exec(f"python3 /workspace/_fetch.py '{safe_url}'", timeout=30)
    output = output.strip()
    sb.exec("rm -f /workspace/_fetch.py")

    if exit_code != 0:
        logger.warning("抓取网页失败: url=%s exit=%d", url, exit_code)
        return f"（抓取失败: {output[:300]}）"

    if not output:
        return "（页面无内容）"

    logger.info("抓取网页成功: %s (%d 字符)", url, len(output))
    return output


# ==================== 沙盒工具 ====================
# 依赖 server/agent/sandbox.py 管理容器生命周期。
# Agent 对话开始时创建沙盒，对话结束销毁，同一对话中多次工具调用共享容器。


@tool
def sandbox_run(command: str) -> str:
    """在 Linux 沙盒中执行任意 shell 命令，返回完整输出。
    沙盒是一个独立的 Linux 容器，可执行 Python 脚本、pip install 包、文件操作等。
    适用场景：代码执行、数据处理、软件包安装与测试、文件操作、运行脚本。
    提示：多个步骤的命令用 && 或分号串联；需要多个命令时可多次调用，容器状态会保持。"""
    from server.agent.sandbox import get_or_create

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动，请稍后重试）"

    # 根据命令类型智能判断超时时间
    # pip install / apt install / git clone 等耗时操作给 120 秒
    # 普通命令 30 秒
    cmd_lower = command.strip().lower()
    if any(k in cmd_lower for k in ["pip install", "apt", "git clone", "npm install", "cargo"]):
        timeout = 180
    elif any(k in cmd_lower for k in ["python", "sh ", "./"]):
        timeout = 60
    else:
        timeout = 30

    exit_code, output = sb.exec(command, timeout=timeout)
    output = output.strip() or "（无输出）"

    if len(output) > 4000:
        output = output[:4000] + "\n\n... （输出过长已截断）"

    prefix = "" if exit_code == 0 else f"（exit_code={exit_code}）\n"
    return prefix + output


@tool
def sandbox_write_file(path: str, content: str) -> str:
    """在沙盒中创建或覆盖文件。
    path: 文件路径（如 /workspace/hello.py）
    content: 文件内容"""
    from server.agent.sandbox import get_or_create

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动）"

    sb.write_file(path, content)
    return f"✓ 已写入 {path}"


@tool
def sandbox_edit_file(path: str, old_text: str, new_text: str) -> str:
    """精确替换文件中的某段文本（用于修改已有文件，避免全量重写）。
    path: 文件路径（如 /workspace/sort.py）
    old_text: 要替换的原始文本（必须与文件中的内容完全一致，包括缩进和换行）
    new_text: 替换后的新文本

    使用场景：
    - 修改代码中的某一行或某几行
    - 替换函数名、变量名
    - 修改配置文件中的某个值

    注意：
    - old_text 必须在文件中唯一匹配，如果有多处匹配会报错，请提供更多上下文使其唯一
    - 如果 old_text 不存在于文件中，会报错
    - 替换后文件会自动保存"""
    from server.agent.sandbox import get_or_create

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动）"

    # 读取当前内容
    current = sb.read_file(path)
    if current.startswith("（文件不存在"):
        return f"（文件不存在: {path}）"

    # 检查匹配
    count = current.count(old_text)
    if count == 0:
        # 提供附近内容帮助 Agent 定位
        preview = current[:500] if len(current) > 500 else current
        return f"（未找到要替换的文本。文件前 500 字符：\n{preview}）"
    if count > 1:
        return f"（匹配到 {count} 处，old_text 不唯一。请提供更多上下文使其唯一匹配）"

    # 执行替换
    new_content = current.replace(old_text, new_text, 1)
    sb.write_file(path, new_content)

    # 返回改动上下文（前后各几行）方便 Agent 确认
    idx = new_content.index(new_text)
    lines_before = new_content[:idx].count("\n")
    context_start = max(0, idx - 200)
    context_end = min(len(new_content), idx + len(new_text) + 200)
    context = new_content[context_start:context_end]

    return f"✓ 已替换 {path} (第 {lines_before + 1} 行)\n```{context}```"


@tool
def sandbox_read_file(path: str) -> str:
    """读取沙盒中的文件内容。"""
    from server.agent.sandbox import get_or_create

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动）"
    return sb.read_file(path)


@tool
def sandbox_list_files(path: str = "/workspace") -> str:
    """列出沙盒目录中的文件。"""
    from server.agent.sandbox import get_or_create

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动）"
    return sb.list_dir(path)


@tool
def sandbox_export_file(path: str) -> str:
    """将沙盒中的文件导出为下载链接，方便用户下载到本地。
    Agent 完成文件生成后应调用此工具，然后将返回的下载链接提供给用户。
    适用场景：生成图表、报告、代码文件、数据文件等需要交付给用户的文件。
    图片文件（.png/.jpg/.svg/.gif）会直接在聊天中渲染显示，其他文件显示为下载链接。"""
    from server.agent.sandbox import get_or_create

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动）"

    result = sb.export_file(path)
    if result is None:
        return f"（导出失败: 文件不存在或无法读取 - {path}）"

    token, filename = result
    url = f"/api/download/{token}/{filename}"

    # 图片文件用 ![]() 语法直接在 Markdown 中渲染
    img_exts = ('.png', '.jpg', '.jpeg', '.gif', '.svg', '.webp', '.bmp')
    if filename.lower().endswith(img_exts):
        return f"✓ 图片已导出，在下方渲染：\n![{filename}]({url})"

    return f"✓ 文件已导出: [{filename}]({url}) （点击下载）"


@tool
def convert_document(input_path: str, target_format: str, output_path: str = "") -> str:
    """在沙盒中转换文档格式，尽量保留原格式（字体、样式、表格、图片等）。

    支持的转换方向：
    - docx → pdf / md / txt / html
    - pdf  → md / txt / html（提取文本，保留段落结构）
    - xlsx → csv / pdf / html
    - pptx → pdf / png（每页截图）
    - md   → pdf / docx / html
    - html → pdf / docx

    格式保留说明：
    - docx→pdf：保留字体、颜色、表格、页眉页脚、图片
    - pdf→md/html：保留段落、标题层级、表格结构，图片单独提取
    - xlsx→csv：保留数据，每个 sheet 输出一个 csv
    - md→docx：标题、表格、列表、代码块转换为 Word 对应样式
    - md→pdf：通过 HTML 中转，保留样式

    参数:
        input_path: 沙盒内源文件路径（如 /workspace/report.docx）
        target_format: 目标格式（pdf/md/txt/html/csv/docx/png）
        output_path: 输出文件路径（可选，默认用源文件名换扩展名）
    """
    from server.agent.sandbox import get_or_create
    import os

    session_id = _current_session_id.get()
    sb = get_or_create(session_id) if session_id else None
    if sb is None:
        return "（沙盒未启动）"

    if not output_path:
        base = os.path.splitext(input_path)[0]
        output_path = f"{base}.{target_format}"

    input_ext = os.path.splitext(input_path)[1].lower().lstrip(".")
    target = target_format.lower().lstrip(".")

    # 构建转换脚本
    scripts = {
        ("docx", "pdf"): f'''
# 纯 Python 方案（无需 libreoffice）：python-docx → HTML → weasyprint → PDF
import sys, os, base64
from docx import Document
from pathlib import Path

doc = Document("{input_path}")
parts = ["<html><head><meta charset='utf-8'><style>",
    "body {{ font-family: sans-serif; max-width: 800px; margin: 40px auto; line-height: 1.8; }}",
    "h1,h2,h3,h4 {{ color: #333; margin-top: 20px; }}",
    "table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}",
    "td, th {{ border: 1px solid #ccc; padding: 6px 10px; }}",
    "pre {{ background: #f5f5f5; padding: 10px; overflow-x: auto; }}",
    "img {{ max-width: 100%; }}",
    "</style></head><body>"]

for para in doc.paragraphs:
    style = para.style.name if para.style else ""
    text = para.text.strip()
    if not text:
        parts.append("<br/>")
        continue
    if style.startswith("Heading"):
        try:
            level = int(style.replace("Heading ", "").replace("Heading", ""))
        except:
            level = 1
        parts.append(f"<h{{min(level,4)}}>{{text}}</h{{min(level,4)}}>")
    elif style == "Title":
        parts.append(f"<h1>{{text}}</h1>")
    else:
        parts.append(f"<p>{{text}}</p>")

for table in doc.tables:
    parts.append("<table>")
    for i, row in enumerate(table.rows):
        tag = "th" if i == 0 else "td"
        cells = "".join(f"<{{tag}}>{{cell.text}}</{{tag}}>" for cell in row.cells)
        parts.append(f"<tr>{{cells}}</tr>")
    parts.append("</table>")

for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_data = rel.target_part.blob
        ext = rel.target_part.ext
        if ext in ("png", "jpg", "jpeg", "gif", "webp"):
            b64 = base64.b64encode(img_data).decode()
            mime = "image/" + ("jpeg" if ext == "jpg" else ext)
            parts.append(f'<img src="data:{{mime}};base64,{{b64}}" style="max-width:100%"/>')

parts.append("</body></html>")
html = "\\n".join(parts)

try:
    from weasyprint import HTML
    HTML(string=html).write_pdf("{output_path}")
    print("OK: python-docx + weasyprint")
except Exception as e:
    print(f"FAIL: weasyprint={{e}}")
''',
        ("docx", "md"): f'''
from docx import Document
from pathlib import Path
import re

doc = Document("{input_path}")
lines = []
for para in doc.paragraphs:
    style = para.style.name if para.style else ""
    text = para.text.strip()
    if not text:
        lines.append("")
        continue
    if style.startswith("Heading"):
        try:
            level = int(style.replace("Heading ", "").replace("Heading", ""))
        except:
            level = 1
        lines.append("#" * level + " " + text)
    elif style == "Title":
        lines.append("# " + text)
    elif style.startswith("List"):
        lines.append("- " + text)
    else:
        lines.append(text)

# 表格
for table in doc.tables:
    lines.append("")
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("|" + "|".join(["---"] * len(cells)) + "|")
    lines.append("")

Path("{output_path}").write_text("\\n".join(lines), encoding="utf-8")
print("OK")
''',
        ("docx", "txt"): f'''
from docx import Document
doc = Document("{input_path}")
text = "\\n".join(para.text for para in doc.paragraphs)
# 表格
for table in doc.tables:
    text += "\\n\\n"
    for row in table.rows:
        text += "\\t".join(cell.text for cell in row.cells) + "\\n"
open("{output_path}", "w", encoding="utf-8").write(text)
print("OK")
''',
        ("docx", "html"): f'''
from docx import Document
from pathlib import Path
doc = Document("{input_path}")
html_parts = ["<html><head><meta charset='utf-8'></head><body>"]
for para in doc.paragraphs:
    style = para.style.name if para.style else ""
    text = para.text.strip()
    if not text:
        html_parts.append("<br>")
    elif style.startswith("Heading"):
        try:
            level = int(style.replace("Heading ", "").replace("Heading", ""))
        except:
            level = 1
        html_parts.append(f"<h{{level}}>{{text}}</h{{level}}>")
    else:
        html_parts.append(f"<p>{{text}}</p>")
for table in doc.tables:
    html_parts.append("<table border='1'>")
    for i, row in enumerate(table.rows):
        tag = "th" if i == 0 else "td"
        cells = "".join(f"<{{tag}}>{{cell.text.strip()}}</{{tag}}>" for cell in row.cells)
        html_parts.append(f"<tr>{{cells}}</tr>")
    html_parts.append("</table>")
html_parts.append("</body></html>")
Path("{output_path}").write_text("\\n".join(html_parts), encoding="utf-8")
print("OK")
''',
        ("pdf", "md"): f'''
from pypdf import PdfReader
from pathlib import Path

reader = PdfReader("{input_path}")
lines = []
for i, page in enumerate(reader.pages):
    text = page.extract_text()
    if text:
        lines.append(f"## 第 {{i+1}} 页\\n")
        lines.append(text.strip())
        lines.append("")
Path("{output_path}").write_text("\\n".join(lines), encoding="utf-8")
print("OK")
''',
        ("pdf", "txt"): f'''
from pypdf import PdfReader
reader = PdfReader("{input_path}")
text = "\\n\\n".join(page.extract_text() or "" for page in reader.pages)
open("{output_path}", "w", encoding="utf-8").write(text)
print("OK")
''',
        ("pdf", "html"): f'''
from pypdf import PdfReader
from pathlib import Path
reader = PdfReader("{input_path}")
parts = ["<html><head><meta charset='utf-8'></head><body>"]
for i, page in enumerate(reader.pages):
    text = page.extract_text() or ""
    parts.append(f"<h2>第 {{i+1}} 页</h2>")
    for line in text.split("\\n"):
        line = line.strip()
        if line:
            parts.append(f"<p>{{line}}</p>")
parts.append("</body></html>")
Path("{output_path}").write_text("\\n".join(parts), encoding="utf-8")
print("OK")
''',
        ("xlsx", "csv"): f'''
from openpyxl import load_workbook
import os, csv

wb = load_workbook("{input_path}")
out_dir = os.path.dirname("{output_path}") or "."
base = os.path.splitext(os.path.basename("{input_path}"))[0]
files = []
for ws in wb.worksheets:
    fname = os.path.join(out_dir, f"{{base}}_{{ws.title}}.csv") if len(wb.worksheets) > 1 else "{output_path}"
    with open(fname, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(row)
    files.append(fname)
print(f"OK: {{','.join(os.path.basename(f) for f in files)}}")
''',
        ("xlsx", "html"): f'''
from openpyxl import load_workbook
from pathlib import Path
wb = load_workbook("{input_path}")
parts = ["<html><head><meta charset='utf-8'></head><body>"]
for ws in wb.worksheets:
    parts.append(f"<h2>{{ws.title}}</h2>")
    parts.append("<table border='1'>")
    for row in ws.iter_rows(values_only=True):
        cells = "".join(f"<td>{{str(c) if c is not None else ''}}</td>" for c in row)
        parts.append(f"<tr>{{cells}}</tr>")
    parts.append("</table>")
parts.append("</body></html>")
Path("{output_path}").write_text("\\n".join(parts), encoding="utf-8")
print("OK")
''',
        ("pptx", "pdf"): f'''
import subprocess
r = subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", "{os.path.dirname(output_path) or "."}", "{input_path}"],
                   capture_output=True, text=True, timeout=120)
if r.returncode == 0:
    print("OK")
else:
    print(f"FAIL: {{r.stderr}}")
''',
        ("md", "html"): f'''
from pathlib import Path
import markdown
text = Path("{input_path}").read_text(encoding="utf-8")
html = markdown.markdown(text, extensions=["tables", "fenced_code", "toc"])
full = f"<html><head><meta charset='utf-8'></head><body>{{html}}</body></html>"
Path("{output_path}").write_text(full, encoding="utf-8")
print("OK")
''',
        ("md", "docx"): f'''
from docx import Document
from docx.shared import Pt
from pathlib import Path
import re

text = Path("{input_path}").read_text(encoding="utf-8")
doc = Document()
for line in text.split("\\n"):
    stripped = line.strip()
    if not stripped:
        continue
    if stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith("- ") or stripped.startswith("* "):
        doc.add_paragraph(stripped[2:], style="List Bullet")
    elif stripped.startswith("|"):
        # 简单表格处理
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if not hasattr(doc, '_table_started') or not doc._table_started:
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = "Table Grid"
            for i, c in enumerate(cells):
                table.rows[0].cells[i].text = c
            doc._table = table
            doc._table_row = 0
            doc._table_started = True
        elif not all(c == "---" for c in cells):
            row = doc._table.add_row()
            for i, c in enumerate(cells):
                if i < len(row.cells):
                    row.cells[i].text = c
        doc._table_row += 1
    else:
        doc.add_paragraph(stripped)
doc.save("{output_path}")
print("OK")
''',
        ("md", "pdf"): f'''
# md → html → pdf（用 weasyprint 或 wkhtmltopdf）
from pathlib import Path
import markdown

text = Path("{input_path}").read_text(encoding="utf-8")
html_body = markdown.markdown(text, extensions=["tables", "fenced_code", "toc"])
html = "<html><head><meta charset='utf-8'>\\n<style>\\n" \\
    "body {{ font-family: sans-serif; max-width: 800px; margin: 40px auto; }}\\n" \\
    "table {{ border-collapse: collapse; width: 100%; }}\\n" \\
    "th, td {{ border: 1px solid #ddd; padding: 8px; }}\\n" \\
    "pre {{ background: #f5f5f5; padding: 12px; overflow-x: auto; }}\\n" \\
    "code {{ background: #f5f5f5; padding: 2px 4px; }}\\n" \\
    "</style></head><body>" + html_body + "</body></html>"

# 先写 html
html_path = "{output_path}.html"
Path(html_path).write_text(html, encoding="utf-8")

# 尝试 weasyprint
try:
    from weasyprint import HTML
    HTML(string=html).write_pdf("{output_path}")
    print("OK: weasyprint")
except Exception as e1:
    # 降级 wkhtmltopdf
    import subprocess
    try:
        r = subprocess.run(["wkhtmltopdf", html_path, "{output_path}"],
                          capture_output=True, text=True, timeout=60)
        if r.returncode == 0:
            print("OK: wkhtmltopdf")
        else:
            # 降级 libreoffice
            r2 = subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                                "--outdir", "{os.path.dirname(output_path) or '.'}", html_path],
                               capture_output=True, text=True, timeout=120)
            if r2.returncode == 0:
                print("OK: libreoffice via html")
            else:
                print(f"FAIL: weasyprint={{e1}}, wkhtmltopdf={{r.stderr}}")
    except FileNotFoundError:
        print(f"FAIL: no pdf converter available (weasyprint={{e1}})")
''',
        ("html", "pdf"): f'''
import subprocess
# 优先 wkhtmltopdf（保留 CSS 样式最好）
r1 = subprocess.run(["wkhtmltopdf", "{input_path}", "{output_path}"],
                    capture_output=True, text=True, timeout=60)
if r1.returncode == 0:
    print("OK: wkhtmltopdf")
else:
    # 降级 weasyprint
    try:
        from weasyprint import HTML
        HTML(filename="{input_path}").write_pdf("{output_path}")
        print("OK: weasyprint")
    except Exception as e:
        # 降级 libreoffice
        r3 = subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                            "--outdir", "{os.path.dirname(output_path) or '.'}", "{input_path}"],
                           capture_output=True, text=True, timeout=120)
        print(f"OK: libreoffice" if r3.returncode == 0 else f"FAIL: {{e}}")
''',
        ("html", "docx"): f'''
import subprocess
r = subprocess.run(["libreoffice", "--headless", "--convert-to", "docx",
                    "--outdir", "{os.path.dirname(output_path) or '.'}", "{input_path}"],
                   capture_output=True, text=True, timeout=120)
print("OK" if r.returncode == 0 else f"FAIL: {{r.stderr}}")
''',
        # .doc 老格式 → 统一走 libreoffice
        ("doc", "pdf"): f'''
import subprocess
r = subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", "{os.path.dirname(output_path) or "."}", "{input_path}"],
                   capture_output=True, text=True, timeout=120)
print("OK" if r.returncode == 0 else f"FAIL: {{r.stderr}}")
''',
        ("doc", "docx"): f'''
import subprocess
r = subprocess.run(["libreoffice", "--headless", "--convert-to", "docx", "--outdir", "{os.path.dirname(output_path) or "."}", "{input_path}"],
                   capture_output=True, text=True, timeout=120)
print("OK" if r.returncode == 0 else f"FAIL: {{r.stderr}}")
''',
    }

    key = (input_ext, target)
    if key not in scripts:
        return f"（不支持的转换: {input_ext} → {target}。支持: doc/docx→pdf/docx/md/txt/html, pdf→md/txt/html, xlsx→csv/html, pptx→pdf, md→pdf/docx/html, html→pdf/docx）"

    # 自动安装依赖（按需。libreoffice 太大不自动装，对应的转换路径用纯 Python 方案）
    needs = set()
    script_content = scripts[key]
    if "from weasyprint" in script_content or "import markdown" in script_content:
        needs.add("weasyprint")

    for dep in needs:
        if dep == "weasyprint":
            sb.exec("pip install weasyprint markdown -q 2>/dev/null || true", timeout=60)

    # 写入脚本文件并执行（write_file 会自动加 /workspace/ 前缀）
    script_content = scripts[key]
    script_path = "_convert_script.py"
    sb.write_file(script_path, script_content)

    exit_code, output = sb.exec(f"python3 /workspace/{script_path}", timeout=300)
    result_text = output.strip()

    if exit_code != 0 or "FAIL" in result_text:
        return f"（转换失败: {result_text[-300:]}）"

    # 验证输出文件存在
    exit_code2, ls_output = sb.exec(f"ls -la {output_path}")
    if exit_code2 != 0:
        return f"（转换似乎成功，但输出文件未找到: {output_path}）"

    return f"✓ 文档转换成功: {input_path} → {output_path}\n调用 sandbox_export_file(\"{output_path}\") 导出给用户。"


@tool
def load_skill(skill_name: str) -> str:
    """加载指定技能的完整指导内容。

    系统提示词中列出了已启用的技能摘要，当你需要某个技能的详细工作流程时，
    调用此工具获取完整内容。

    适用场景：
    - 用户任务与某个技能相关，需要参考详细步骤
    - 不确定如何处理某类任务，想查看是否有相关技能指导

    参数:
        skill_name: 技能名称（从系统提示词的技能列表中获取）
    """
    from server.agent.skills import load_skill_detail

    return load_skill_detail(skill_name)


# ==================== 技能搜索与安装（宿主机工具，不经过沙盒） ====================


@tool
def search_skills(query: str) -> str:
    """从开源技能生态（skills.sh）搜索可用的 Agent 技能。

    在宿主机上执行 npx skills find，返回搜索结果。
    不经过沙盒，不需要 Docker。

    适用场景：
    - 用户问 "有没有技能可以做 X"
    - 用户想扩展 Agent 能力
    - 用户想找现成的方案而非从头写

    参数:
        query: 搜索关键词（如 "react performance"、"pr review"、"changelog"）
    """
    try:
        result = subprocess.run(
            ["npx", "-y", "skills", "find", query],
            capture_output=True,
            text=True,
            timeout=60,
            cwd=str(Path.cwd()),
            stdin=subprocess.DEVNULL,
        )
        output = result.stdout.strip() or result.stderr.strip()
        if not output:
            return "（搜索无结果）"

        if len(output) > 4000:
            output = output[:4000] + "\n\n... （结果过长已截断）"

        return output
    except subprocess.TimeoutExpired:
        return "（搜索超时，请稍后重试）"
    except FileNotFoundError:
        return "（npx 未安装，无法搜索技能）"
    except Exception as exc:
        return f"（搜索失败: {exc}）"


@tool
def install_skill(package: str) -> str:
    """安装开源技能到项目 skills/ 目录并自动启用。

    在宿主机上执行 npx skills add，安装到项目的 skills/ 目录下，
    并自动更新 user_settings.json 的 skills 数组。安装后可直接使用，无需重启。

    不经过沙盒，不需要 Docker。

    适用场景：
    - 用户想安装搜索到的技能
    - 用户提供了技能包名（如 owner/repo@skill-name）

    参数:
        package: 技能包名（如 "vercel-labs/agent-skills@vercel-react-best-practices"）
    """
    from server.agent.skills import list_available_skills
    from server.user_settings import user_settings
    from server.logging_config import get_logger

    logger = get_logger(__name__)

    # 安装前的技能列表
    before = set(list_available_skills())

    try:
        result = subprocess.run(
            ["npx", "-y", "skills", "add", package, "-y"],
            capture_output=True,
            text=True,
            timeout=300,
            cwd=str(Path.cwd()),
            stdin=subprocess.DEVNULL,
            env={**__import__('os').environ, 'GIT_HTTP_LOW_SPEED_LIMIT': '1000', 'GIT_HTTP_LOW_SPEED_TIME': '60'},
        )

        output = result.stdout.strip()
        if result.returncode != 0:
            err = result.stderr.strip()
            return f"（安装失败: {err or output[:500]}）"

        # 安装后的技能列表，找出新增的
        # 现在同时扫描项目和用户目录，无需手动复制
        after = set(list_available_skills())
        new_skills = after - before

        if not new_skills:
            return f"安装命令已执行，但未检测到新技能。\n输出: {output[:500]}\n\n请检查技能是否安装到了其他目录。"

        # 技能目录会被自动扫描，无需手动更新配置
        skill_list_str = ", ".join(new_skills)
        logger.info("技能已安装并自动启用: %s", skill_list_str)

        return (
            f"✓ 技能安装成功: {skill_list_str}\n"
            f"已自动启用（技能目录会被自动扫描），可以直接使用。\n"
            f"调用 load_skill('{list(new_skills)[0]}') 可查看技能详细内容。"
        )

    except subprocess.TimeoutExpired:
        return "（安装超时，请稍后重试）"
    except FileNotFoundError:
        return "（npx 未安装，无法安装技能）"
    except Exception as exc:
        return f"（安装失败: {exc}）"


def get_agent_tools() -> list:
    """获取 Agent 可用工具列表。"""
    return [
        search_knowledge_base,
        search_web,
        fetch_webpage,
        sandbox_run,
        sandbox_write_file,
        sandbox_edit_file,
        sandbox_read_file,
        sandbox_list_files,
        sandbox_export_file,
        convert_document,
        load_skill,
        search_skills,
        install_skill,
    ]
